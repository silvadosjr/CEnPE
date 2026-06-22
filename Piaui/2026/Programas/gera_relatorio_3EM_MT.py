from __future__ import annotations

import tempfile
import os
from datetime import date
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SERIES = int(os.environ.get("REPORT_SERIES", "3"))
SERIES_TAG = f"{SERIES}EM"
SERIES_ORDINAL = "2ª" if SERIES == 2 else "3ª"
DISCIPLINE = os.environ.get("REPORT_DISCIPLINE", "MT").upper()
DISCIPLINE_NAME = "Matemática" if DISCIPLINE == "MT" else "Língua Portuguesa"
PROFILE = f"{SERIES_TAG}_{DISCIPLINE}"
OUTPUT_DIR = ROOT / "Relatorios" / PROFILE
OUTPUT = OUTPUT_DIR / f"Relatorio_Tecnico_{PROFILE}_TCM_TRI_refinado.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
CAUTION = "FFF4CE"
RISK = "FDE9E7"
WHITE = "FFFFFF"


def fmt_num(value: float, decimals: int = 1) -> str:
    return f"{value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_int(value: int) -> str:
    return f"{value:,}".replace(",", ".")


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill, border=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, color=MUTED, italic=True)
    return p


def add_picture_with_alt(doc, path, width, alt_text):
    shape = doc.add_picture(str(path), width=width)
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return shape


def add_source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Fonte: {text}")
    set_font(run, size=8.5, color=MUTED)
    return p


def add_table(doc, headers, rows, widths_dxa, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(label))
        set_font(run, size=9, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_font(run, size=9.2, color="222222")
    set_table_geometry(table, widths_dxa)
    return table


def read_data():
    data = {}
    for key, roman, suffix in (("S1", "I", "S1"), ("S2", "II", "S2")):
        tcm = ROOT / f"ResultadosTCM_{PROFILE}/Simulado_{roman}"
        tri = ROOT / f"ResultadosTRI_{PROFILE}/Simulado_{roman}"
        scores = pd.read_excel(tcm / "EscoresBrutos_Aluno.xlsx")
        measures = pd.read_excel(tcm / "MedidasTCM_Item.xlsx")
        theta = pd.read_excel(tri / f"ResumoTheta_{PROFILE}_{suffix}.xlsx", index_col=0).iloc[:, 0]
        theta_ic = pd.read_excel(tri / f"ResumoThetaIC_{PROFILE}_{suffix}.xlsx", index_col=0)
        items = pd.read_excel(tri / f"EstItens_{PROFILE}_{suffix}.xlsx")
        dist = pd.read_excel(tri / f"DistriAlunosClasse{DISCIPLINE}{SERIES}serie_{suffix}.xlsx", index_col=0)
        alt_choice = pd.read_excel(tcm / "mDificNRDF.xlsx", index_col=0)
        alt_disc = pd.read_excel(tcm / "mDiscNRDF.xlsx", index_col=0)
        alt_pbis = pd.read_excel(tcm / "mcpBisNRDF.xlsx", index_col=0)
        data[key] = {
            "scores": scores,
            "measures": measures,
            "theta": theta,
            "theta_ic": theta_ic,
            "items": items,
            "dist": dist,
            "alt_choice": alt_choice,
            "alt_disc": alt_disc,
            "alt_pbis": alt_pbis,
        }
    return data


def make_charts(data, workdir):
    colors = {"S1": "#2E74B5", "S2": "#D9822B"}
    try:
        font = ImageFont.truetype("arial.ttf", 20)
        small = ImageFont.truetype("arial.ttf", 14)
        tiny = ImageFont.truetype("arial.ttf", 11)
    except OSError:
        font = small = tiny = ImageFont.load_default()

    def bar_panel(draw, box, title, values, vmax, base_color, problem_indices=None, reference=None, shade=None):
        x0, y0, x1, y1 = box
        draw.text((x0, y0), title, font=font, fill="#203748")
        top, bottom, left, right = y0 + 36, y1 - 34, x0 + 48, x1 - 12
        if shade:
            lo, hi = shade
            sy1 = bottom - (hi / vmax) * (bottom - top)
            sy0 = bottom - (lo / vmax) * (bottom - top)
            draw.rectangle((left, sy1, right, sy0), fill="#E8EEF5")
        for tick in range(0, 5):
            value = vmax * tick / 4
            yy = bottom - (value / vmax) * (bottom - top)
            draw.line((left, yy, right, yy), fill="#D9DEE5", width=1)
            draw.text((x0, yy - 7), f"{value:.2f}".replace(".", ","), font=tiny, fill="#5F6B76")
        if reference is not None:
            yy = bottom - (reference / vmax) * (bottom - top)
            for xx in range(left, right, 10):
                draw.line((xx, yy, min(xx + 5, right), yy), fill="#607D8B", width=2)
        draw.line((left, top, left, bottom), fill="#5F6B76", width=1)
        draw.line((left, bottom, right, bottom), fill="#5F6B76", width=1)
        gap = (right - left) / len(values)
        bw = max(4, int(gap * 0.68))
        for idx, value in enumerate(values, start=1):
            cx = left + (idx - 0.5) * gap
            fill = "#B3261E" if problem_indices and idx in problem_indices else base_color
            if float(value) >= 0:
                yy = bottom - (float(value) / vmax) * (bottom - top)
                draw.rectangle((cx - bw / 2, yy, cx + bw / 2, bottom), fill=fill)
            else:
                draw.rectangle((cx - bw / 2, bottom, cx + bw / 2, bottom + 12), fill=fill)
            if idx % 2 == 0:
                draw.text((cx - 7, bottom + 5), str(idx), font=tiny, fill="#5F6B76")

    img = Image.new("RGB", (1800, 1180), "white")
    draw = ImageDraw.Draw(img)
    panels = [(40, 30, 880, 570), (920, 30, 1760, 570), (40, 610, 880, 1150), (920, 610, 1760, 1150)]
    for col, key in enumerate(("S1", "S2")):
        measures = data[key]["measures"]
        roman = "I" if key == "S1" else "II"
        if PROFILE == "3EM_MT":
            problems = {8} if key == "S1" else {14}
        elif PROFILE == "3EM_LP":
            problems = {1} if key == "S1" else {8, 12, 20}
        elif PROFILE == "2EM_LP":
            problems = {9, 10, 21} if key == "S1" else set()
        else:
            problems = set()
        bar_panel(draw, panels[col], f"Simulado {roman} - proporção de acertos", measures["Dificuldade"].tolist(), 0.90, colors[key], problems, shade=(0.2, 0.4))
        bar_panel(draw, panels[col + 2], f"Simulado {roman} - discriminação", measures["Discriminacao"].tolist(), 0.75, colors[key], problems, reference=0.20)
    tcm_chart = workdir / "tcm_itens.png"
    img.save(tcm_chart, dpi=(180, 180))

    classes = [str(x) for x in data["S1"]["dist"].index]
    classes = [c for c in classes if any(c in data[k]["dist"].index and float(data[k]["dist"].loc[c, "N de alunos"]) > 0 for k in ("S1", "S2"))]
    labels = []
    for c in classes:
        if c == "[350,1e+03)":
            labels.append("350+")
        else:
            labels.append(c.replace("[", "").replace(")", "").replace(",", "–"))
    values = {
        key: [float(data[key]["dist"].loc[c, "Percentual"]) if c in data[key]["dist"].index else 0.0 for c in classes]
        for key in ("S1", "S2")
    }
    img = Image.new("RGB", (1800, 800), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 30), "Distribuição dos estudantes por faixa de proficiência", font=font, fill="#203748")
    left, right, top, bottom = 120, 1740, 100, 680
    vmax = 30
    for tick in range(0, 31, 5):
        yy = bottom - (tick / vmax) * (bottom - top)
        draw.line((left, yy, right, yy), fill="#D9DEE5", width=1)
        draw.text((65, yy - 8), str(tick), font=small, fill="#5F6B76")
    draw.text((20, 370), "Estudantes (%)", font=small, fill="#5F6B76")
    draw.line((left, top, left, bottom), fill="#5F6B76", width=2)
    draw.line((left, bottom, right, bottom), fill="#5F6B76", width=2)
    group = (right - left) / len(labels)
    bw = int(group * 0.32)
    for idx, label in enumerate(labels):
        cx = left + (idx + 0.5) * group
        for offset, key in ((-bw / 2, "S1"), (bw / 2, "S2")):
            value = values[key][idx]
            yy = bottom - (value / vmax) * (bottom - top)
            draw.rectangle((cx + offset - bw / 2, yy, cx + offset + bw / 2, bottom), fill=colors[key])
        draw.text((cx - 35, bottom + 12), label, font=tiny, fill="#5F6B76")
    draw.rectangle((1310, 35, 1340, 55), fill=colors["S1"])
    draw.text((1350, 36), "Simulado I", font=small, fill="#203748")
    draw.rectangle((1500, 35, 1530, 55), fill=colors["S2"])
    draw.text((1540, 36), "Simulado II", font=small, fill="#203748")
    tri_chart = workdir / "distribuicao_theta.png"
    img.save(tri_chart, dpi=(180, 180))
    return tcm_chart, tri_chart


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("PIAUÍ 2026  |  AVALIAÇÃO EDUCACIONAL")
    set_font(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build_report(data, charts):
    doc = Document()
    configure_document(doc)
    tcm_chart, tri_chart = charts

    # Capa: editorial_cover.
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RELATÓRIO TÉCNICO")
    set_font(r, size=11, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Análises TCM e TRI")
    set_font(r, size=30, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{SERIES_ORDINAL} série do Ensino Médio – {DISCIPLINE_NAME}")
    set_font(r, size=15, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Simulados I e II | Piauí, 2026")
    set_font(r, size=11, color=MUTED, bold=True)
    p.paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento-base metodológico: Relatório 3EM-MT 2025")
    set_font(r, size=10.5, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(42)
    months = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    today = date.today()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{today.day} de {months[today.month - 1]} de {today.year}")
    set_font(r, size=10, color=MUTED)
    doc.add_page_break()

    s1, s2 = data["S1"], data["S2"]
    n1, n2 = len(s1["scores"]), len(s2["scores"])
    mean1, mean2 = s1["scores"].iloc[:, 1].mean(), s2["scores"].iloc[:, 1].mean()
    theta1, theta2 = float(s1["theta"].loc["media"]), float(s2["theta"].loc["media"])

    doc.add_heading("1. Síntese executiva", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(p, "EDF4FB", BLUE)
    if PROFILE == "3EM_MT":
        lead = (
            "Leitura principal. As duas aplicações apresentaram provas exigentes para a população avaliada. "
            "No Simulado I, o item 8 mostrou discriminação praticamente nula; no Simulado II, o mesmo ocorreu com o item 14. "
            "O item 8 foi mantido na nova calibração TRI, apesar da ressalva; o item 14 permaneceu excluído, em coerência com os indícios de conflito de gabarito."
        )
    elif PROFILE == "3EM_LP":
        lead = (
            "Leitura principal. O Simulado I apresentou um item com funcionamento fraco (item 1), mas manteve os 26 itens na TRI. "
            "No Simulado II, os itens 8, 12 e 20 tiveram discriminação e correlação ponto-bisserial abaixo de 0,20, com sinais de conflito de gabarito; "
            "os três foram excluídos da calibração TRI."
        )
    elif PROFILE == "2EM_LP":
        lead = (
            "Leitura principal. No Simulado I, os itens 9, 10 e 21 apresentaram discriminação e correlação ponto-bisserial abaixo de 0,20. "
            "O item 21 foi excluído da TRI; o item 9 permaneceu na calibração, mas combinou sinal de conflito de gabarito, baixa discriminação e dificuldade TRI extrema. "
            "O Simulado II não apresentou itens abaixo dos limiares diagnósticos da TCM."
        )
    else:
        lead = (
            "Leitura principal. As duas aplicações apresentaram provas exigentes, mas todos os 26 itens mantiveram discriminação e correlação "
            "ponto-bisserial acima de 0,20. Nenhum item foi excluído da calibração TRI. Os itens 24, no Simulado I, e 22, no Simulado II, "
            "foram os mais difíceis e devem ser acompanhados, sem evidência isolada de funcionamento anômalo."
        )
    r = p.add_run(lead)
    set_font(r, size=10.5, color=NAVY, bold=True)

    tri_counts = [len(s1["items"]), len(s2["items"])]
    tcm_item_sets = [set(s1["measures"]["Item"].astype(int)), set(s2["measures"]["Item"].astype(int))]
    tri_item_sets = [set(s1["items"]["Item"].astype(int)), set(s2["items"]["Item"].astype(int))]
    exclusions = []
    for tcm_set, tri_set in zip(tcm_item_sets, tri_item_sets):
        missing = sorted(tcm_set - tri_set)
        exclusions.append(", ".join(str(x) for x in missing) if missing else "Nenhum")
    rows = [
        ["Simulado I", fmt_int(n1), f"26 / {tri_counts[0]}", f"{fmt_num(mean1, 2)} / 26", f"{fmt_num(100 * mean1 / 26, 1)}%", fmt_num(theta1, 1), exclusions[0]],
        ["Simulado II", fmt_int(n2), f"26 / {tri_counts[1]}", f"{fmt_num(mean2, 2)} / 26", f"{fmt_num(100 * mean2 / 26, 1)}%", fmt_num(theta2, 1), exclusions[1]],
    ]
    add_table(
        doc,
        ["Aplicação", "N", "Itens TCM / TRI", "Escore médio", "% do total", "Proficiência média", "Exclusão TRI"],
        rows,
        [1300, 1000, 1300, 1350, 1150, 1900, 1360],
        [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6,
    )
    add_source(doc, f"ResumoEscores, ResumoTheta, EscoresBrutos_Aluno e EstItens da {PROFILE.replace('_', '-')}.")
    doc.add_paragraph(
        "A diferença descritiva entre as médias foi de "
        f"{fmt_num(mean2 - mean1, 2)} ponto no escore bruto e {fmt_num(theta2 - theta1, 1)} pontos na escala de proficiência. "
        "Esses valores não devem ser interpretados isoladamente como crescimento longitudinal: as bases têm tamanhos distintos e o relatório não estabelece que os estudantes, a cobertura curricular e as condições de aplicação sejam equivalentes."
    )

    doc.add_heading("2. Introdução e abordagem analítica", level=1)
    doc.add_paragraph(
        f"Foram analisadas as respostas dos estudantes avaliados em {DISCIPLINE_NAME} na {SERIES_ORDINAL} série do Ensino Médio. "
        f"A primeira aplicação reuniu {fmt_int(n1)} registros e a segunda, {fmt_int(n2)}. "
        "Em cada simulado, a TCM considerou os 26 itens apresentados. A TRI utilizou um modelo logístico de três parâmetros (3PL), "
        f"com itens conhecidos fixados para posicionamento na escala SAEB e {len(s1['items'])} itens no Simulado I e {len(s2['items'])} no Simulado II."
    )
    doc.add_paragraph(
        "A organização deste documento preserva a linha do Relatório Técnico 3EM-MT de 2025: a TCM é usada de forma descritiva para examinar escores, "
        "proporção de acertos, discriminação e funcionamento das alternativas; a TRI é usada para equalização, estimação dos parâmetros dos itens e descrição "
        "da proficiência na escala transformada. A ampliação para dois simulados exige, contudo, que os contrastes entre aplicações sejam apresentados com cautela."
    )
    doc.add_heading("Teoria Clássica de Medidas", level=2)
    doc.add_paragraph(
        "Na TCM, o desempenho é sintetizado pelo escore bruto e por estatísticas calculadas diretamente das respostas observadas. "
        "A proporção de acertos descreve a dificuldade empírica do item; o índice de discriminação e as correlações bisserial e ponto-bisserial informam em que medida "
        "o item separa estudantes de maior e menor desempenho. A leitura por alternativa permite detectar distratores positivamente associados ao escore e possíveis conflitos de gabarito [1, 2]."
    )
    doc.add_heading("Modelo TRI e equalização", level=2)
    doc.add_paragraph(
        "Foi ajustado o modelo logístico de três parâmetros para um único grupo (3PL). A estimação dos itens empregou distribuições a priori e o algoritmo EM; "
        "os traços latentes individuais foram estimados pelo método da esperança a posteriori (EAP) [3–5]."
    )
    for equation in (
        "Yᵢⱼ | θⱼ, ζᵢ  ~  Bernoulli(pᵢⱼ)",
        "pᵢⱼ = cᵢ + (1 - cᵢ) / {1 + exp[-aᵢ(θⱼ - bᵢ)]}",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(equation)
        set_font(r, name="Cambria Math", size=11, color=NAVY)
    doc.add_paragraph(
        "Nesse modelo, aᵢ representa a discriminação, bᵢ a dificuldade, cᵢ a probabilidade de acerto casual e θⱼ a proficiência do estudante. "
        "Os itens previamente calibrados foram fixados durante a estimação. Ao final, as proficiências e as dificuldades foram expressas na escala de referência com transformação 250 + 50θ."
    )
    doc.add_heading("Convenções de leitura", level=2)
    for text in (
        "Na TCM, a coluna denominada “dificuldade” corresponde à proporção de acertos: valores menores indicam itens mais difíceis.",
        "Itens com discriminação ou correlação ponto-bisserial abaixo de 0,20 foram tratados como sinais de alerta, em consonância com as linhas de referência utilizadas nos scripts.",
        "Na TRI, a proficiência individual foi estimada por EAP e transformada para a escala com centro 250 e desvio 50; os resumos apresentados descrevem a distribuição das estimativas individuais.",
        "Parâmetros extremos e estatísticas de ajuste devem ser interpretados em conjunto com o conteúdo do item, o gabarito e as curvas características, não como critérios automáticos de descarte.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3. Resultados da TCM", level=1)
    rows = []
    for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
        d = data[key]
        scores = d["scores"].iloc[:, 1]
        m = d["measures"]
        rows.append([
            label,
            fmt_num(scores.mean(), 2),
            fmt_num(scores.std(), 2),
            fmt_num(scores.median(), 0),
            fmt_num(m["Dificuldade"].mean(), 3),
            str(int(((m["Dificuldade"] >= 0.20) & (m["Dificuldade"] < 0.40)).sum())),
            str(int((m["Dificuldade"] >= 0.40).sum())),
        ])
    add_table(
        doc,
        ["Aplicação", "Média", "DP", "Mediana", "Prop. média", "Itens 0,20–0,40", "Itens ≥0,40"],
        rows,
        [1450, 1100, 1000, 1100, 1500, 1650, 1560],
        [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6,
    )
    add_source(doc, "ResumoEscores e MedidasTCM_Item.")
    band_s1 = int(((s1["measures"]["Dificuldade"] >= 0.20) & (s1["measures"]["Dificuldade"] < 0.40)).sum())
    band_s2 = int(((s2["measures"]["Dificuldade"] >= 0.20) & (s2["measures"]["Dificuldade"] < 0.40)).sum())
    ge40_s1 = int((s1["measures"]["Dificuldade"] >= 0.40).sum())
    ge40_s2 = int((s2["measures"]["Dificuldade"] >= 0.40).sum())
    ge60_s1 = int((s1["measures"]["Dificuldade"] >= 0.60).sum())
    ge60_s2 = int((s2["measures"]["Dificuldade"] >= 0.60).sum())
    ease_sentence = (
        "Nenhum item alcançou 0,60 de acerto."
        if ge60_s1 == 0 and ge60_s2 == 0
        else f"A proporção de acertos alcançou pelo menos 0,60 em {ge60_s1} itens do Simulado I e {ge60_s2} do Simulado II."
    )
    if DISCIPLINE == "LP":
        if PROFILE == "3EM_LP":
            diagnostic_sentence = "Apesar do perfil global menos difícil, o Simulado II reuniu dois itens abaixo de 0,20 de acerto e três itens com funcionamento discriminativo fraco."
        else:
            diagnostic_sentence = "No Simulado I, um item ficou abaixo de 0,20 de acerto e três apresentaram funcionamento discriminativo fraco; no Simulado II, nenhum item ficou abaixo dos limiares de 0,20."
        tcm_overview = (
            f"A maior parte dos itens apresentou proporção de acertos a partir de 0,40: {ge40_s1} itens no Simulado I e {ge40_s2} no Simulado II. "
            f"{ease_sentence} {diagnostic_sentence}"
        )
    else:
        tcm_overview = (
            f"A concentração dos itens entre 0,20 e 0,40 de acerto — {band_s1} itens no Simulado I e {band_s2} no Simulado II — caracteriza instrumentos predominantemente difíceis. "
            f"{ease_sentence} O Simulado II apresentou {ge40_s2} itens com pelo menos 0,40 de acerto, contra {ge40_s1} no Simulado I."
        )
    doc.add_paragraph(tcm_overview)
    add_picture_with_alt(
        doc,
        tcm_chart,
        Inches(6.35),
        "Quatro gráficos de barras mostram a proporção de acertos e a discriminação dos 26 itens nos Simulados I e II."
        + (" Os itens de atenção aparecem destacados em vermelho." if PROFILE in {"3EM_MT", "3EM_LP", "2EM_LP"} else ""),
    )
    caption = "Figura 1. Proporção de acertos e discriminação por item."
    if PROFILE in {"3EM_MT", "3EM_LP", "2EM_LP"}:
        caption += " Itens de atenção destacados em vermelho."
    add_caption(doc, caption)
    add_source(doc, "MedidasTCM_Item, Simulados I e II.")

    if PROFILE == "3EM_MT":
        doc.add_heading("Itens com comportamento atípico", level=2)
        add_table(
            doc,
            ["Aplicação", "Item", "Prop. de acertos", "Discriminação", "Corr. ponto-bisserial", "Encaminhamento"],
            [
                ["Simulado I", "8", "0,291", "0,063", "0,022", "Mantido na TRI; interpretar com ressalva"],
                ["Simulado II", "14", "0,253", "0,090", "0,034", "Excluído da TRI; revisar gabarito"],
            ],
            [1250, 700, 1450, 1300, 1650, 3010],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        )
        add_source(doc, "MedidasTCM_Item e Impressões gerais sobre os resultados.")
        doc.add_heading("Diagnóstico das alternativas", level=2)
        alt_rows = []
        for key, label, item_number in (("S1", "Simulado I", 8), ("S2", "Simulado II", 14)):
            d = data[key]
            row_choice = d["alt_choice"].iloc[item_number - 1]
            row_disc = d["alt_disc"].iloc[item_number - 1]
            row_pbis = d["alt_pbis"].iloc[item_number - 1]
            for alternative in ("A", "B", "C", "D", "E"):
                alt_rows.append([
                    label, str(item_number), alternative, fmt_num(row_choice[alternative], 1) + "%",
                    fmt_num(row_disc[alternative], 3), fmt_num(row_pbis[alternative], 3),
                    "Gabarito" if alternative == row_choice["Gabarito"] else "Distrator",
                ])
        add_table(
            doc,
            ["Aplicação", "Item", "Alternativa", "% escolha", "Discriminação", "Ponto-bisserial", "Papel"],
            alt_rows,
            [1250, 650, 1100, 1150, 1450, 1600, 2160],
            [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6,
        )
        add_source(doc, "mDificNRDF, mDiscNRDF e mcpBisNRDF.")
        doc.add_paragraph(
            "Nos dois itens, a alternativa A foi mais escolhida do que a alternativa D, registrada como gabarito, e apresentou associação positiva mais forte com o desempenho. "
            "No item 8 do Simulado I, A reuniu 36,4% das escolhas e ponto-bisserial 0,180, enquanto D reuniu 29,1% e ponto-bisserial 0,022. "
            "No item 14 do Simulado II, A reuniu 38,8% e ponto-bisserial 0,139, contra 25,3% e 0,034 para D. Esse padrão reforça a hipótese de conflito de gabarito."
        )
        p = doc.add_paragraph()
        set_paragraph_shading(p, RISK, "B3261E")
        r = p.add_run(
            "A baixa discriminação, acompanhada de correlação ponto-bisserial próxima de zero, indica que o acerto nesses itens quase não acompanhou o desempenho geral. "
            "Esse padrão é compatível com gabarito problemático, ambiguidade ou funcionamento anômalo. O item 8 foi mantido na TRI como item fixado, mas seus resultados devem continuar acompanhados da ressalva técnica."
        )
        set_font(r, size=10.5, color="6B1D16")
    elif PROFILE == "3EM_LP":
        doc.add_heading("Itens com comportamento atípico", level=2)
        issue_rows = []
        for key, label, item_numbers in (("S1", "Simulado I", [1]), ("S2", "Simulado II", [8, 12, 20])):
            measures = data[key]["measures"]
            for item_number in item_numbers:
                row = measures.loc[measures["Item"].astype(int) == item_number].iloc[0]
                issue_rows.append([
                    label, str(item_number), fmt_num(row["Dificuldade"], 3),
                    fmt_num(row["Discriminacao"], 3), fmt_num(row["CPBisserial"], 3),
                    "Excluído da TRI; revisar gabarito" if key == "S2" else "Mantido na TRI; revisar funcionamento",
                ])
        add_table(
            doc,
            ["Aplicação", "Item", "Prop. de acertos", "Discriminação", "Corr. ponto-bisserial", "Encaminhamento"],
            issue_rows,
            [1250, 700, 1450, 1300, 1650, 3010],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        )
        add_source(doc, "MedidasTCM_Item e Impressões gerais sobre os resultados.")
        doc.add_paragraph(
            "O item 1 do Simulado I teve discriminação 0,154 e ponto-bisserial 0,131, justificando revisão, embora tenha permanecido na TRI. "
            "No Simulado II, os itens 8, 12 e 20 apresentaram discriminação baixa ou negativa e foram retirados da calibração TRI."
        )
        doc.add_heading("Diagnóstico das alternativas no Simulado II", level=2)
        alt_rows = []
        for item_number in (8, 12, 20):
            row_choice = s2["alt_choice"].iloc[item_number - 1]
            row_disc = s2["alt_disc"].iloc[item_number - 1]
            row_pbis = s2["alt_pbis"].iloc[item_number - 1]
            for alternative in ("A", "B", "C", "D", "E"):
                alt_rows.append([
                    str(item_number), alternative, fmt_num(row_choice[alternative], 1) + "%",
                    fmt_num(row_disc[alternative], 3), fmt_num(row_pbis[alternative], 3),
                    "Gabarito" if alternative == row_choice["Gabarito"] else "Distrator",
                ])
        add_table(
            doc,
            ["Item", "Alternativa", "% escolha", "Discriminação", "Ponto-bisserial", "Papel"],
            alt_rows,
            [800, 1150, 1400, 1700, 1900, 2410],
            [WD_ALIGN_PARAGRAPH.CENTER] * 6,
        )
        add_source(doc, "mDificNRDF, mDiscNRDF e mcpBisNRDF, Simulado II.")
        doc.add_paragraph(
            "No item 8, o gabarito B reuniu 11,9% das escolhas e ponto-bisserial -0,049, enquanto o distrator E reuniu 58,4% e ponto-bisserial 0,362. "
            "No item 20, o gabarito A reuniu 21,1% e ponto-bisserial 0,034, contra 47,1% e 0,280 para o distrator B. "
            "No item 12, o gabarito A teve 16,4% das escolhas e ponto-bisserial 0,099; a alternativa E foi a mais escolhida, com 33,8%, embora também tenha apresentado associação fraca com o desempenho."
        )
        p = doc.add_paragraph()
        set_paragraph_shading(p, RISK, "B3261E")
        r = p.add_run(
            "O padrão dos itens 8 e 20 é fortemente compatível com conflito de gabarito, pois uma alternativa não oficial concentrou mais respostas e discriminou melhor. "
            "O item 12 também requer revisão, mas seu padrão é menos conclusivo e pode envolver ambiguidade, conteúdo ou funcionamento geral fraco."
        )
        set_font(r, size=10.5, color="6B1D16")
    elif PROFILE == "2EM_LP":
        doc.add_heading("Itens com comportamento atípico no Simulado I", level=2)
        issue_rows = []
        for item_number in (9, 10, 21):
            row = s1["measures"].loc[s1["measures"]["Item"].astype(int) == item_number].iloc[0]
            issue_rows.append([
                "Simulado I", str(item_number), fmt_num(row["Dificuldade"], 3),
                fmt_num(row["Discriminacao"], 3), fmt_num(row["CPBisserial"], 3),
                "Excluído da TRI; revisar gabarito" if item_number == 21 else "Mantido na TRI; revisar funcionamento",
            ])
        add_table(
            doc,
            ["Aplicação", "Item", "Prop. de acertos", "Discriminação", "Corr. ponto-bisserial", "Encaminhamento"],
            issue_rows,
            [1250, 700, 1450, 1300, 1650, 3010],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        )
        add_source(doc, "MedidasTCM_Item e Impressões gerais sobre os resultados.")
        doc.add_paragraph(
            "Os itens 9, 10 e 21 ficaram abaixo de 0,20 em discriminação e correlação ponto-bisserial. O item 21 foi retirado da TRI. "
            "O item 9 permaneceu na calibração, mas apresentou dificuldade estimada de 587,3 pontos e discriminação TRI de 0,349, reforçando a necessidade de revisão."
        )
        doc.add_heading("Diagnóstico das alternativas no Simulado I", level=2)
        alt_rows = []
        for item_number in (9, 10, 21):
            row_choice = s1["alt_choice"].iloc[item_number - 1]
            row_disc = s1["alt_disc"].iloc[item_number - 1]
            row_pbis = s1["alt_pbis"].iloc[item_number - 1]
            for alternative in ("A", "B", "C", "D", "E"):
                alt_rows.append([
                    str(item_number), alternative, fmt_num(row_choice[alternative], 1) + "%",
                    fmt_num(row_disc[alternative], 3), fmt_num(row_pbis[alternative], 3),
                    "Gabarito" if alternative == row_choice["Gabarito"] else "Distrator",
                ])
        add_table(
            doc,
            ["Item", "Alternativa", "% escolha", "Discriminação", "Ponto-bisserial", "Papel"],
            alt_rows,
            [800, 1150, 1400, 1700, 1900, 2410],
            [WD_ALIGN_PARAGRAPH.CENTER] * 6,
        )
        add_source(doc, "mDificNRDF, mDiscNRDF e mcpBisNRDF, Simulado I.")
        doc.add_paragraph(
            "No item 9, o gabarito E reuniu 11,9% das escolhas e ponto-bisserial 0,140, enquanto o distrator A reuniu 45,5% e ponto-bisserial 0,258. "
            "No item 21, o gabarito E teve 23,8% das escolhas e ponto-bisserial 0,065; B foi a alternativa mais escolhida (29,7%) e D apresentou a maior associação positiva com o desempenho (0,132). "
            "No item 10, o gabarito B foi a alternativa mais escolhida e a única com associação positiva relevante, mas o ponto-bisserial permaneceu baixo (0,162)."
        )
        p = doc.add_paragraph()
        set_paragraph_shading(p, RISK, "B3261E")
        r = p.add_run(
            "O item 9 apresenta o sinal mais forte de conflito de gabarito. O item 21 também mostra funcionamento inconsistente, com respostas distribuídas entre alternativas concorrentes, "
            "o que sustenta sua exclusão da TRI. O item 10 parece sobretudo pouco discriminativo, sem evidência clara de troca de gabarito."
        )
        set_font(r, size=10.5, color="6B1D16")
    else:
        doc.add_heading("Itens de maior dificuldade", level=2)
        attention_rows = []
        for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
            for _, row in data[key]["measures"].nsmallest(3, "Dificuldade").iterrows():
                attention_rows.append([
                    label, str(row["Item"]).zfill(3), fmt_num(row["Dificuldade"], 3),
                    fmt_num(row["Discriminacao"], 3), fmt_num(row["CPBisserial"], 3),
                ])
        add_table(
            doc,
            ["Aplicação", "Item", "Prop. de acertos", "Discriminação", "Corr. ponto-bisserial"],
            attention_rows,
            [1600, 900, 2100, 2100, 2660],
            [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
        )
        add_source(doc, "MedidasTCM_Item.")
        doc.add_paragraph(
            "O item 24 do Simulado I e o item 22 do Simulado II foram os únicos com proporção de acertos abaixo de 0,20. "
            "Apesar da dificuldade elevada, ambos preservaram discriminação acima de 0,20 e correlação ponto-bisserial positiva, não configurando, isoladamente, conflito de gabarito. "
            "Nenhum item apresentou discriminação ou correlação ponto-bisserial abaixo do limiar diagnóstico de 0,20."
        )

    doc.add_heading("4. Resultados da TRI", level=1)
    rows = []
    for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
        d = data[key]
        th = d["theta"]
        ic = d["theta_ic"].loc["SAEB"]
        rows.append([
            label,
            fmt_num(th.loc["media"], 1),
            f"{fmt_num(ic['IC_Inf'], 1)} a {fmt_num(ic['IC_Sup'], 1)}",
            fmt_num(th.loc["dp"], 1),
            fmt_num(th.loc["med."], 1),
            f"{fmt_num(th.loc['1o Q'], 1)} a {fmt_num(th.loc['3oQ'], 1)}",
            str(int((d["items"]["SAEB"] == "Sim").sum())),
        ])
    add_table(
        doc,
        ["Aplicação", "Média", "IC 95% da média", "DP EAP", "Mediana", "Intervalo interquartil", "Itens fixados"],
        rows,
        [1300, 950, 1700, 1000, 1050, 1900, 1460],
        [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6,
    )
    add_source(doc, "ResumoTheta e ResumoThetaIC.")
    add_picture_with_alt(
        doc,
        tri_chart,
        Inches(6.35),
        "Gráfico de barras agrupadas compara os percentuais de estudantes dos Simulados I e II em oito faixas de proficiência, de 175 a 350 pontos ou mais.",
    )
    add_caption(doc, "Figura 2. Distribuição percentual dos estudantes por faixa de proficiência.")
    add_source(doc, f"DistriAlunosClasse{DISCIPLINE}{SERIES}serie, Simulados I e II.")
    doc.add_paragraph(
        f"No Simulado I, a proficiência média foi {fmt_num(theta1, 1)}, com mediana {fmt_num(s1['theta'].loc['med.'], 1)} e assimetria positiva ({fmt_num(s1['theta'].loc['ca'], 2)}). "
        f"No Simulado II, a média foi {fmt_num(theta2, 1)}, a mediana {fmt_num(s2['theta'].loc['med.'], 1)} e a assimetria foi {fmt_num(s2['theta'].loc['ca'], 2)}. "
        f"A proporção na faixa de 350 pontos ou mais foi {fmt_num(s1['dist'].iloc[-1]['Percentual'], 1)}% na primeira aplicação e {fmt_num(s2['dist'].iloc[-1]['Percentual'], 1)}% na segunda."
    )

    doc.add_heading("Parâmetros dos itens", level=2)
    param_rows = []
    param_summary = []
    for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
        items = data[key]["items"]
        checks = (
            ("bSAEB > 400", items[items["bSAEB"] > 400]),
            ("a < 0,60", items[items["a"] < 0.60]),
            ("a > 4,00", items[items["a"] > 4.00]),
            ("c > 0,25", items[items["c"] > 0.25]),
        )
        for signal, subset in checks:
            if not subset.empty:
                item_text = ", ".join(str(int(x)) for x in subset["Item"].tolist())
                param_rows.append([label, signal, item_text])
        n_b = int((items["bSAEB"] > 400).sum())
        n_a = int((items["a"] < 0.60).sum())
        n_c = int((items["c"] > 0.25).sum())
        param_summary.append(
            f"{label}: {n_b} {'item' if n_b == 1 else 'itens'} com bSAEB acima de 400, "
            f"{n_a} {'item' if n_a == 1 else 'itens'} com a abaixo de 0,60 e {n_c} {'item' if n_c == 1 else 'itens'} com c acima de 0,25"
        )
    doc.add_paragraph(
        "; ".join(param_summary) + ". Esses valores são sinais diagnósticos e devem ser conferidos em conjunto com erros-padrão, ajuste e curvas do item."
    )
    add_table(
        doc,
        ["Aplicação", "Sinal para revisão", "Itens"],
        param_rows,
        [1450, 2300, 5610],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_source(doc, f"EstItens_{PROFILE}. Os pontos de corte são diagnósticos e não implicam exclusão automática.")

    doc.add_heading("5. Leitura comparativa entre as aplicações", level=1)
    if PROFILE == "3EM_LP":
        comparison_text = (
            "O Simulado II apresentou escore bruto médio e proporção de acertos superiores, mas a proficiência média TRI foi ligeiramente inferior à do Simulado I. "
            "Também houve menor dispersão das proficiências e redução da proporção de estudantes na faixa de 350 pontos ou mais. "
            "A divergência entre escore bruto e escala equalizada reforça que as formas têm dificuldades e composições distintas; além disso, três itens do Simulado II foram retirados da TRI. "
            "O contraste deve, portanto, permanecer descritivo e não ser interpretado isoladamente como evolução ou queda de aprendizagem."
        )
    elif PROFILE == "2EM_LP":
        comparison_text = (
            "O Simulado II apresentou escore bruto médio e proficiência média superiores, acompanhados de maior dispersão das proficiências e aumento da proporção de estudantes na faixa de 350 pontos ou mais. "
            "A TCM mostra proporção média de acertos semelhante entre as formas, com eliminação dos sinais fracos observados em três itens do Simulado I. "
            "Como os conjuntos de estudantes e itens não são integralmente equivalentes e um item foi retirado da TRI do Simulado I, o contraste deve permanecer descritivo."
        )
    else:
        comparison_text = (
            "O Simulado II apresentou escore bruto médio e proficiência média superiores, menor dispersão relativa das proficiências e maior concentração nas faixas a partir de 300 pontos. "
            "A TCM também mostra uma forma ligeiramente menos difícil, com maior proporção média de acertos. Em contrapartida, os conjuntos de itens e de estudantes não são integralmente documentados aqui como equivalentes; "
            "por isso, o contraste deve permanecer descritivo."
        )
    doc.add_paragraph(comparison_text)
    add_table(
        doc,
        ["Indicador", "Simulado I", "Simulado II", "Diferença descritiva"],
        [
            ["N de estudantes", fmt_int(n1), fmt_int(n2), f"+{fmt_int(n2 - n1)}"],
            ["Escore bruto médio", fmt_num(mean1, 2), fmt_num(mean2, 2), f"+{fmt_num(mean2 - mean1, 2)}"],
            ["Proporção média de acertos", fmt_num(s1["measures"]["Dificuldade"].mean(), 3), fmt_num(s2["measures"]["Dificuldade"].mean(), 3), f"+{fmt_num(s2['measures']['Dificuldade'].mean() - s1['measures']['Dificuldade'].mean(), 3)}"],
            ["Proficiência média", fmt_num(theta1, 1), fmt_num(theta2, 1), f"+{fmt_num(theta2 - theta1, 1)}"],
            ["DP das proficiências EAP", fmt_num(s1["theta"].loc["dp"], 1), fmt_num(s2["theta"].loc["dp"], 1), fmt_num(s2["theta"].loc["dp"] - s1["theta"].loc["dp"], 1)],
            ["Estudantes com 350+", f"{fmt_num(s1['dist'].iloc[-1]['Percentual'], 1)}%", f"{fmt_num(s2['dist'].iloc[-1]['Percentual'], 1)}%", f"+{fmt_num(s2['dist'].iloc[-1]['Percentual'] - s1['dist'].iloc[-1]['Percentual'], 1)} p.p."],
        ],
        [2700, 1850, 1850, 2960],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_source(doc, f"Resultados TCM e TRI da {PROFILE.replace('_', '-')}. Diferenças sem ajuste por composição ou pareamento.")

    doc.add_heading("6. Recomendações", level=1)
    if PROFILE == "3EM_MT":
        recommendations = [
            "Manter documentada a ressalva do item 8 do Simulado I, que permaneceu na TRI como item fixado, e evitar interpretação isolada de seu desempenho.",
            "Revisar formalmente o gabarito e o conteúdo do item 14 do Simulado II, que permaneceu excluído da TRI.",
            "Examinar os distratores desses dois itens por grupo de escore, verificando se uma alternativa não oficial foi progressivamente mais escolhida pelos estudantes de maior desempenho.",
            "Realizar leitura de ajuste TRI item a item, combinando infit/outfit, testes X²/G²/PV-Q1, curvas características e erros-padrão antes de aprovar itens com parâmetros extremos.",
            "Para comparar os simulados como medida de evolução, estabelecer base pareada de estudantes ou aplicar ponderação/controle de composição, além de confirmar a estabilidade dos itens de ligação.",
            "Produzir uma ficha de decisão por item com status: aprovado, revisar, excluir da proficiência ou reservar para nova aplicação.",
        ]
    elif PROFILE == "3EM_LP":
        recommendations = [
            "Revisar formalmente os gabaritos e o conteúdo dos itens 8, 12 e 20 do Simulado II, documentando a decisão final antes de qualquer reutilização.",
            "Revisar também o item 1 do Simulado I, que permaneceu na TRI, mas apresentou discriminação e correlação ponto-bisserial abaixo de 0,20.",
            "Examinar os distratores por grupo de escore, especialmente E no item 8 e B no item 20, para confirmar se a alternativa não oficial foi progressivamente mais escolhida pelos estudantes de maior desempenho.",
            "Realizar leitura de ajuste TRI item a item, combinando infit/outfit, testes X²/G²/PV-Q1, curvas características e erros-padrão antes de aprovar itens com parâmetros extremos.",
            "Para comparar os simulados como medida de evolução, estabelecer base pareada de estudantes ou aplicar ponderação/controle de composição, além de confirmar a estabilidade dos itens de ligação.",
        ]
    elif PROFILE == "2EM_LP":
        recommendations = [
            "Revisar formalmente os gabaritos e o conteúdo dos itens 9 e 21 do Simulado I, documentando a decisão final antes de qualquer reutilização.",
            "Revisar o item 10 do Simulado I quanto à clareza, ao alinhamento curricular e à qualidade dos distratores, sem presumir conflito de gabarito.",
            "Submeter o item 9 a nova análise ou excluí-lo de interpretações de proficiência, pois combinou baixa discriminação TCM, dificuldade TRI extrema e alternativa concorrente mais informativa.",
            "Realizar leitura de ajuste TRI item a item, combinando infit/outfit, testes X²/G²/PV-Q1, curvas características e erros-padrão antes de aprovar itens com parâmetros extremos.",
            "Para comparar os simulados como medida de evolução, estabelecer base pareada de estudantes ou aplicar ponderação/controle de composição, além de confirmar a estabilidade dos itens de ligação.",
        ]
    else:
        recommendations = [
            "Manter os 26 itens nas bases analíticas, pois nenhum apresentou discriminação ou correlação ponto-bisserial abaixo de 0,20 na TCM.",
            "Revisar pedagogicamente os itens 24 (Simulado I) e 22 (Simulado II), os mais difíceis, verificando alinhamento curricular e clareza sem presumir erro de gabarito.",
            "Realizar leitura de ajuste TRI item a item, combinando infit/outfit, testes X²/G²/PV-Q1, curvas características e erros-padrão, com atenção aos itens de dificuldade ou acerto casual elevados.",
            "Para comparar os simulados como medida de evolução, estabelecer base pareada de estudantes ou aplicar ponderação/controle de composição, além de confirmar a estabilidade dos itens de ligação.",
            "Produzir uma ficha de decisão por item com status: aprovado, revisar, reservar para nova aplicação ou excluir da proficiência.",
        ]
    for item in recommendations:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("7. Referências e arquivos-fonte", level=1)
    references = [
        "[1] Primi, R. (2012). Psicometria: fundamentos matemáticos da Teoria Clássica dos Testes. Avaliação Psicológica, 11, 297–307.",
        "[2] Wu, M.; Tam, H. P.; Jen, T.-H. (2017). Educational Measurement for Applied Researchers: Theory into Practice. Springer.",
        "[3] Andrade, D. F.; Tavares, H. R.; Valle, R. C. (2000). Teoria da Resposta ao Item: conceitos e aplicações. ABE.",
        "[4] Baker, F. B.; Kim, S.-H. (2004). Item Response Theory: Parameter Estimation Techniques. 2. ed. Marcel Dekker.",
        "[5] Azevedo, C. L. N. (2003). Métodos de estimação na Teoria da Resposta ao Item. Dissertação de Mestrado, IME-USP.",
        "[6] R Core Team (2020). R: A Language and Environment for Statistical Computing. R Foundation for Statistical Computing.",
        "[7] Relatório Técnico 3EM-MT 2025. Avaliação Diagnóstica em Matemática - 3ª Série EM - Piauí.",
    ]
    for ref in references:
        doc.add_paragraph(ref)
    doc.add_heading("Arquivos de resultados utilizados", level=2)
    sources = [
        f"ResultadosTCM_{PROFILE}/Simulado_I e Simulado_II: ResumoEscores.xlsx, MedidasTCM_Item.xlsx, PropAcertosGrupo.xlsx e PropEscolhaAlternativaGrupo.xlsx.",
        f"ResultadosTRI_{PROFILE}/Simulado_I e Simulado_II: ResumoTheta, ResumoThetaIC, EstItens, DistriAlunosClasse e BaseRespTheta.",
        f"Programas/TCM_{PROFILE}_S1.R, TCM_{PROFILE}_S2.R, Fit_3PL_Fix_{PROFILE}_S1.R e Fit_3PL_Fix_{PROFILE}_S2.R.",
    ]
    if PROFILE in {"3EM_MT", "3EM_LP", "2EM_LP"}:
        sources.append("Impressões gerais..odt, com as ressalvas técnicas registradas pela equipe.")
    for src in sources:
        doc.add_paragraph(src, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Apêndice A. Medidas TCM por item", level=1)
    for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
        doc.add_heading(label, level=2)
        rows = []
        for _, row in data[key]["measures"].iterrows():
            rows.append([
                str(row["Item"]).zfill(3),
                fmt_num(row["Dificuldade"], 3),
                fmt_num(row["Discriminacao"], 3),
                fmt_num(row["CBisserial"], 3),
                fmt_num(row["CPBisserial"], 3),
            ])
        add_table(
            doc,
            ["Item", "Proporção de acertos", "Discriminação", "Corr. bisserial", "Corr. ponto-bisserial"],
            rows,
            [900, 2200, 1800, 2000, 2460],
            [WD_ALIGN_PARAGRAPH.CENTER] * 5,
        )
        add_source(doc, f"ResultadosTCM_{PROFILE}/{label.replace(' ', '_')}/MedidasTCM_Item.xlsx.")
        if key == "S1":
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Apêndice B. Parâmetros TRI por item", level=1)
    for key, label in (("S1", "Simulado I"), ("S2", "Simulado II")):
        doc.add_heading(label, level=2)
        rows = []
        for _, row in data[key]["items"].iterrows():
            rows.append([
                str(row["Item"]).zfill(3),
                fmt_num(row["a"], 3),
                fmt_num(row["b"], 3),
                fmt_num(row["bSAEB"], 1),
                fmt_num(row["c"], 3),
                "Sim" if row["SAEB"] == "Sim" else "Não",
            ])
        add_table(
            doc,
            ["Item", "a", "b", "b (escala 250,50)", "c", "Fixado"],
            rows,
            [900, 1300, 1300, 2500, 1400, 1960],
            [WD_ALIGN_PARAGRAPH.CENTER] * 6,
        )
        add_source(doc, f"ResultadosTRI_{PROFILE}/{label.replace(' ', '_')}/EstItens_{PROFILE}_{key}.xlsx.")
        if key == "S1":
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Apêndice C. Distribuição por faixa de proficiência", level=1)
    classes = [str(x) for x in data["S1"]["dist"].index]
    rows = []
    for classe in classes:
        row = [classe]
        for key in ("S1", "S2"):
            d = data[key]["dist"]
            if classe in d.index:
                row.extend([fmt_int(int(d.loc[classe, "N de alunos"])), fmt_num(float(d.loc[classe, "Percentual"]), 3) + "%"])
            else:
                row.extend(["0", "0,000%"])
        rows.append(row)
    add_table(
        doc,
        ["Faixa", "N - Sim. I", "% - Sim. I", "N - Sim. II", "% - Sim. II"],
        rows,
        [2200, 1700, 1700, 1700, 2060],
        [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_source(doc, f"DistriAlunosClasse{DISCIPLINE}{SERIES}serie, Simulados I e II.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main():
    data = read_data()
    with tempfile.TemporaryDirectory(prefix="relatorio_3em_mt_") as tmp:
        charts = make_charts(data, Path(tmp))
        output = build_report(data, charts)
    print(output)


if __name__ == "__main__":
    main()
